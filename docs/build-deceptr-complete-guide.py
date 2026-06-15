from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "DECEPTR_Guide_Complet_Installation_Configuration.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 98, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WARN = "FFF2CC"
RISK = "FCE4D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05

    small = styles.add_style("SmallNote", 1)
    small.font.name = "Calibri"
    small.font.size = Pt(9.5)
    small.font.color.rgb = MUTED
    small.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "DECEPTR v1 MVP - Guide complet d'installation et configuration"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    set_page_number(footer)


def add_para(doc, text="", style=None, bold=False, color=None, size=None, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.left_indent = Inches(0.18)
        run = p.add_run(line.rstrip())
        set_run_font(run, name="Consolas", size=8.5, color=RGBColor(35, 45, 60))


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    add_para(doc, "PROJET CYBERDECEPTION", bold=True, color=BLUE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "DECEPTR v1 MVP", bold=True, color=NAVY, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Guide complet - Architecture - Installation manuelle - Configuration - Tests - Exploitation",
        color=MUTED,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["Champ", "Valeur"],
        [
            ["Institution", "Chambre des Representants du Maroc - DSI / Cybersecurite"],
            ["Projet", "Plateforme de cyberdeception DECEPTR v1 MVP"],
            ["Sujet", "Honeypot Cowrie T-Pot, pipeline ELK, alerting et dashboard SOC"],
            ["Stack", "Docker, T-Pot Cowrie, Filebeat, Logstash, Redis, Python, Elasticsearch, MySQL, MinIO, Kibana, FastAPI, ElastAlert"],
            ["Date", "Juin 2026"],
            ["Dossier final", r"D:\assir\Ismagi\PFA\DECEPTR-FINAL"],
        ],
        widths=[1.55, 4.95],
    )
    add_callout(
        doc,
        "Ce document couvre",
        "Les erreurs connues, l'architecture complete, l'installation manuelle, la configuration manuelle de chaque service, les commandes de demarrage, les tests bout-en-bout, les dashboards, la maintenance et la checklist de validation.",
    )
    page_break(doc)


def add_static_toc(doc):
    doc.add_heading("Sommaire", level=1)
    items = [
        "1. Erreurs et problemes connus",
        "2. Architecture globale du projet",
        "3. Prerequis et preparation de la machine",
        "4. Construire DECEPTR de zero",
        "5. Installation manuelle depuis le dossier final",
        "6. Configuration manuelle des composants",
        "7. Demarrage manuel de toute l'architecture",
        "8. Configuration manuelle de Kibana et des dashboards",
        "9. Tests utilisateur et validation bout-en-bout",
        "10. Scripts d'automatisation du projet",
        "11. Exploitation SOC et rapport final",
        "12. Maintenance, sauvegarde et nettoyage Docker",
        "13. Checklist complete",
        "14. Conclusion",
    ]
    numbered(doc, items)
    page_break(doc)


def section_errors(doc):
    doc.add_heading("1. Erreurs et problemes connus - Realite du terrain", level=1)
    add_para(
        doc,
        "Avant d'installer DECEPTR, il faut connaitre les pannes qui arrivent le plus souvent dans un lab Docker/ELK/T-Pot. Ces points correspondent aux problemes rencontres pendant la mise en route du projet.",
    )
    add_table(
        doc,
        ["Erreur", "Gravite", "Symptome", "Cause", "Solution"],
        [
            ["Pull Docker timeout", "Critique", "docker compose echoue sur auth.docker.io TLS handshake timeout", "Connexion lente ou Docker Hub temporairement indisponible", "Relancer le pull, garder Docker Desktop ouvert, eviter VPN instable"],
            ["Kibana vide", "Majeur", "Kibana s'ouvre mais ne montre aucun dashboard", "Aucun data view cree ou aucun evenement indexe", "Creer data views cowrie-* et deceptr-events-* puis generer un evenement test"],
            ["Port 22 occupe", "Critique", "Cowrie ne demarre pas", "OpenSSH Windows ou autre service utilise deja le port 22", "Arreter le service SSH Windows ou changer le mapping en lab"],
            ["Login dashboard ne reagit pas", "Majeur", "Mot de passe saisi mais rien ne se passe", "API indisponible, CORS/URL API incorrecte ou MySQL non pret", "Verifier http://127.0.0.1:8000/health et les logs deceptr-api"],
            ["Filebeat refuse la config", "Majeur", "Exiting: error loading config file", "Permissions strictes Filebeat", "Lancer avec --strict.perms=false"],
            ["TLS Filebeat vers Logstash KO", "Critique", "x509 ou handshake failed", "CA absente ou mauvais certificat", "Copier elk/certs/ca.crt dans tpot-runtime/deceptr/certs"],
            ["Elasticsearch unhealthy", "Critique", "Kibana/API/Pipeline attendent Elasticsearch", "Memoire insuffisante ou volume corrompu", "Donner 4-6 GB RAM a Docker Desktop; supprimer seulement les volumes DECEPTR si necessaire"],
            ["T-Pot complet trop lourd", "Moyen", "Beaucoup de conteneurs non prevus apparaissent", "Stack T-Pot complet lance au lieu du runtime strict", "Utiliser tpot-runtime avec seulement cowrie et deceptr-tpot-forwarder"],
            ["Pas d'alertes", "Moyen", "Dashboard affiche 0 alertes", "Aucun evenement Cowrie traite", "Lancer le test E2E ou faire une tentative SSH/Telnet"],
        ],
        widths=[1.25, 0.75, 1.35, 1.55, 1.6],
        header_fill=RISK,
    )
    add_callout(
        doc,
        "Regle d'or",
        "Ne pas lancer deux versions du meme DECEPTR en meme temps. Le projet final propre est D:\\assir\\Ismagi\\PFA\\DECEPTR-FINAL. Les autres projets Docker ne doivent pas etre supprimes.",
        fill=WARN,
    )


def section_architecture(doc):
    doc.add_heading("2. Architecture globale du projet", level=1)
    add_para(doc, "DECEPTR est une plateforme de cyberdeception. Elle expose un honeypot SSH/Telnet en DMZ, collecte les traces, enrichit les evenements, calcule un risque et presente les alertes au SOC.")
    add_code(
        doc,
        r"""
Internet / Attaquants
        |
        v
Firewall / DMZ
        |
        v
T-Pot Cowrie officiel (SSH 22 / Telnet 23)
        |
        v
Filebeat T-Pot Forwarder -- TLS 1.3 / TCP 5044 -->
        |
        v
DECEPTR Logstash -> Redis Queue -> Pipeline Python
        |
        +--> Elasticsearch: cowrie-* / deceptr-events-*
        +--> MySQL: alerts, iocs, attackers, campaigns, users
        +--> MinIO: malware-samples, downloads, reports, backups
        |
        v
Kibana + FastAPI + Dashboard DECEPTR + ElastAlert 2
""",
    )
    add_table(
        doc,
        ["Couche", "Composants", "Role"],
        [
            ["Detection", "cowrie", "Honeypot SSH/Telnet officiel T-Pot"],
            ["Collecte", "deceptr-tpot-forwarder", "Lecture de cowrie.json et envoi TLS vers Logstash"],
            ["Ingestion", "deceptr-logstash", "Parsing, index brut, push vers Redis"],
            ["Queue", "deceptr-redis", "Tampon entre Logstash et le pipeline"],
            ["Traitement", "deceptr-pipeline", "Collector, normalizer, enricher, correlator, MITRE mapper, risk scorer, alerter"],
            ["Stockage", "Elasticsearch, MySQL, MinIO", "Logs, evenements, donnees metier, objets"],
            ["Visualisation", "Kibana, FastAPI, Dashboard", "Exploration, API REST, interface SOC"],
            ["Alerting", "ElastAlert 2, alerter Python", "Regles, notifications, alertes critiques"],
        ],
        widths=[1.15, 2.0, 3.35],
    )
    doc.add_heading("2.1 Dossier final", level=2)
    add_code(
        doc,
        r"""
D:\assir\Ismagi\PFA\DECEPTR-FINAL
|-- deceptr
|   |-- docker-compose.yml
|   |-- docker-compose.tpot.yml
|   |-- dashboard
|   |-- elk
|   |-- mysql
|   |-- pipeline
|   |-- tpot
|   `-- scripts
|-- tpot-runtime
|   |-- docker-compose.yml
|   |-- .env
|   |-- data\cowrie\log
|   `-- deceptr\filebeat + certs
|-- docs
|-- start.ps1
`-- stop.ps1
""",
    )
    add_callout(doc, "Etat valide actuel", "Le stack final ne contient pas SpiderFoot, ni ancien T-Pot web, ni services hors schema. Les conteneurs attendus sont: cowrie, deceptr-tpot-forwarder, deceptr-logstash, deceptr-redis, deceptr-pipeline, deceptr-elasticsearch, deceptr-mysql, deceptr-minio, deceptr-kibana, deceptr-api, deceptr-dashboard, deceptr-elastalert.")


def section_prereq(doc):
    doc.add_heading("3. Prerequis et preparation de la machine", level=1)
    add_table(
        doc,
        ["Element", "Minimum lab", "Recommande"],
        [
            ["OS", "Windows 10/11 avec WSL2", "Windows 11 + Docker Desktop stable"],
            ["RAM Docker", "4 GB", "6 a 8 GB pour ELK confortable"],
            ["CPU", "2 coeurs", "4 coeurs"],
            ["Disque", "20 GB libres", "50 GB libres"],
            ["Ports libres", "22, 23, 5044, 5601, 8000, 8088, 9000, 9001, 9200", "Memes ports reserves au projet"],
            ["Reseau", "Internet pour pull images", "Connexion stable sans proxy bloquant Docker Hub"],
        ],
        widths=[1.3, 2.6, 2.6],
    )
    doc.add_heading("3.1 Verifications Windows", level=2)
    add_code(
        doc,
        r"""
docker version
docker compose version
wsl --status
netstat -ano | findstr ":22"
netstat -ano | findstr ":23"
netstat -ano | findstr ":5601"
""",
    )
    add_para(doc, "Si le port 22 est occupe par OpenSSH Server Windows, arreter le service avant de lancer Cowrie:")
    add_code(
        doc,
        r"""
Get-Service sshd
Stop-Service sshd
Set-Service sshd -StartupType Manual
""",
    )
    add_callout(doc, "Important securite", "Les mots de passe inclus sont des valeurs de lab. Pour un deploiement reel, changer ELASTIC_PASSWORD, MYSQL_PASSWORD, REDIS_PASSWORD, JWT_SECRET, MINIO_ROOT_PASSWORD et les cles API.", fill=WARN)


def section_install(doc):
    doc.add_heading("5. Installation manuelle depuis le dossier final", level=1)
    doc.add_heading("4.1 Recuperer le dossier final", level=2)
    add_para(doc, "Le dossier final a ete consolide dans un seul repertoire. Pour une installation manuelle sur une autre machine, copier ce dossier complet sans separer deceptr et tpot-runtime.")
    add_code(
        doc,
        r"""
D:\assir\Ismagi\PFA\DECEPTR-FINAL
""",
    )
    doc.add_heading("4.2 Installer Docker Desktop", level=2)
    numbered(
        doc,
        [
            "Installer Docker Desktop pour Windows.",
            "Activer WSL2 backend dans Settings > General.",
            "Dans Settings > Resources, donner au minimum 4 GB RAM.",
            "Demarrer Docker Desktop et attendre que l'icone indique Running.",
            "Tester docker version et docker compose version.",
        ],
    )
    doc.add_heading("4.3 Charger les images manuellement", level=2)
    add_para(doc, "Si on veut verifier image par image avant de lancer compose:")
    add_code(
        doc,
        r"""
docker pull ghcr.io/telekom-security/cowrie:24.04.1
docker pull elastic/filebeat:8.13.0
docker pull logstash:8.13.0
docker pull elasticsearch:8.13.0
docker pull docker.elastic.co/kibana/kibana:8.13.0
docker pull redis:7-alpine
docker pull mysql:8.0
docker pull minio/minio:latest
docker pull nginx:1.27-alpine
docker pull jertel/elastalert2:latest
""",
    )
    doc.add_heading("4.4 Construire les images Python", level=2)
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr
docker compose -f docker-compose.yml -f docker-compose.tpot.yml build api pipeline
""",
    )


def section_build_zero(doc):
    doc.add_heading("4. Construire DECEPTR de zero", level=1)
    add_para(
        doc,
        "Cette section explique comment reconstruire le projet depuis une machine vide, sans utiliser le dossier final deja prepare. Elle sert si on veut refaire le projet pour un rapport, une soutenance ou une nouvelle installation propre.",
    )
    add_callout(
        doc,
        "Principe",
        "On cree d'abord l'arborescence, puis les fichiers Docker Compose, les certificats TLS, la configuration Logstash/Filebeat, la base MySQL, le pipeline Python, l'API, le dashboard, le runtime Cowrie T-Pot et enfin les scripts start/stop.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("4.1 Creer l'arborescence", level=2)
    add_code(
        doc,
        r"""
mkdir D:\assir\Ismagi\PFA\DECEPTR-FROM-ZERO
cd D:\assir\Ismagi\PFA\DECEPTR-FROM-ZERO

mkdir deceptr
mkdir deceptr\elk\certs
mkdir deceptr\elk\logstash\pipeline
mkdir deceptr\elk\logstash\config
mkdir deceptr\elk\elastalert\rules
mkdir deceptr\config\redis
mkdir deceptr\mysql
mkdir deceptr\pipeline\api\routes
mkdir deceptr\pipeline\api\models
mkdir deceptr\dashboard
mkdir deceptr\tpot\filebeat
mkdir deceptr\scripts
mkdir tpot-runtime
mkdir tpot-runtime\data\cowrie\log
mkdir tpot-runtime\deceptr\filebeat
mkdir tpot-runtime\deceptr\certs
mkdir docs
""",
    )

    doc.add_heading("4.2 Creer le fichier .env", level=2)
    add_para(doc, "Dans deceptr\\.env, mettre les secrets de lab. En production, changer toutes ces valeurs.")
    add_code(
        doc,
        r"""
ELASTIC_PASSWORD=deceptr2025
KIBANA_SYSTEM_PASS=deceptr-kibana
REDIS_PASSWORD=deceptr-redis-2026
REDIS_HOST=redis
REDIS_PORT=6379
REDIS_QUEUE=deceptr:events
MYSQL_ROOT_PASSWORD=root2025
MYSQL_PASSWORD=mysql2025
MINIO_ROOT_USER=deceptr_admin
MINIO_ROOT_PASSWORD=MotDePasseMinIO2026!
MINIO_ENDPOINT=minio:9000
MINIO_ACCESS_KEY=deceptr_admin
MINIO_SECRET_KEY=MotDePasseMinIO2026!
MINIO_SECURE=false
JWT_SECRET=deceptr-jwt-change-me-32chars-min
VIRUSTOTAL_KEY=
ABUSEIPDB_KEY=
COWRIE_HOSTNAME=srv-prod-01
SMTP_HOST=
SMTP_PORT=587
SMTP_USER=
SMTP_PASS=
ALERT_TO=security@example.com
""",
    )

    doc.add_heading("4.3 Generer les certificats TLS", level=2)
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FROM-ZERO\deceptr\elk\certs
openssl genrsa -out ca.key 4096
openssl req -x509 -new -nodes -key ca.key -sha256 -days 3650 -out ca.crt -subj "/CN=DECEPTR-CA"
openssl genrsa -out logstash.key 2048
openssl req -new -key logstash.key -out logstash.csr -subj "/CN=logstash"
openssl x509 -req -in logstash.csr -CA ca.crt -CAkey ca.key -CAcreateserial -out logstash.crt -days 825 -sha256
copy ca.crt D:\assir\Ismagi\PFA\DECEPTR-FROM-ZERO\tpot-runtime\deceptr\certs\ca.crt
""",
    )

    doc.add_heading("4.4 Creer docker-compose.yml DECEPTR", level=2)
    add_para(doc, "Le compose principal doit contenir les services suivants. On peut recopier ce fichier depuis le projet final ou le reconstruire manuellement avec cette structure.")
    add_table(
        doc,
        ["Service", "Image/build", "Role", "Ports"],
        [
            ["logstash", "logstash:8.13.0", "Ingestion Beats TLS, parsing, Redis, Elasticsearch", "5044"],
            ["redis", "redis:7-alpine", "Queue d'evenements", "6379"],
            ["elasticsearch", "elasticsearch:8.13.0", "Index cowrie-* et deceptr-events-*", "9200"],
            ["kibana", "docker.elastic.co/kibana/kibana:8.13.0", "Exploration et dashboards", "5601"],
            ["mysql", "mysql:8.0", "Donnees metier", "interne 3306"],
            ["minio", "minio/minio:latest", "Stockage objets", "9000/9001"],
            ["pipeline", "build ./pipeline", "Traitement Python", "interne"],
            ["api", "build ./pipeline", "FastAPI", "8000"],
            ["dashboard", "nginx:1.27-alpine", "Interface SOC statique", "8088"],
            ["elastalert", "jertel/elastalert2:latest", "Regles alerting", "interne"],
        ],
        widths=[1.25, 1.9, 2.25, 1.1],
    )
    add_para(doc, "Creer aussi docker-compose.tpot.yml pour desactiver les services locaux cowrie/filebeat du compose principal:")
    add_code(
        doc,
        r"""
services:
  cowrie:
    profiles: ["local-cowrie"]
  filebeat:
    profiles: ["local-filebeat"]
""",
    )

    doc.add_heading("4.5 Creer le runtime Cowrie T-Pot", level=2)
    add_para(doc, "Dans tpot-runtime\\docker-compose.yml, garder seulement Cowrie officiel T-Pot et Filebeat forwarder.")
    add_code(
        doc,
        r"""
networks:
  cowrie_local:
  deceptr_internal:
    external: true

services:
  cowrie:
    image: ghcr.io/telekom-security/cowrie:24.04.1
    container_name: cowrie
    restart: unless-stopped
    networks: [cowrie_local]
    ports:
      - "22:22"
      - "23:23"
    read_only: true
    volumes:
      - ./data/cowrie/downloads:/home/cowrie/cowrie/dl
      - ./data/cowrie/keys:/home/cowrie/cowrie/etc
      - ./data/cowrie/log:/home/cowrie/cowrie/log
      - ./data/cowrie/log/tty:/home/cowrie/cowrie/log/tty

  deceptr-tpot-forwarder:
    image: elastic/filebeat:8.13.0
    container_name: deceptr-tpot-forwarder
    restart: unless-stopped
    user: root
    command: ["--strict.perms=false"]
    environment:
      - DECEPTR_LOGSTASH_HOST=logstash
      - DECEPTR_LOGSTASH_PORT=5044
    networks:
      - cowrie_local
      - deceptr_internal
    volumes:
      - ./deceptr/filebeat/tpot-cowrie-to-deceptr.yml:/usr/share/filebeat/filebeat.yml:ro
      - ./deceptr/certs/ca.crt:/usr/share/filebeat/certs/ca.crt:ro
      - ./data/cowrie/log:/tpot/cowrie/log:ro
""",
    )

    doc.add_heading("4.6 Creer Filebeat T-Pot", level=2)
    add_para(doc, "Fichier: deceptr\\tpot\\filebeat\\tpot-cowrie-to-deceptr.yml puis copie vers tpot-runtime\\deceptr\\filebeat.")
    add_code(
        doc,
        r"""
filebeat.inputs:
  - type: log
    enabled: true
    paths:
      - /tpot/cowrie/log/cowrie.json
    json.keys_under_root: true
    json.overwrite_keys: true
    json.add_error_key: true
    tags: ["cowrie", "tpot"]
    fields:
      deceptr_source: tpot-cowrie

output.logstash:
  hosts: ["${DECEPTR_LOGSTASH_HOST}:${DECEPTR_LOGSTASH_PORT}"]
  ssl.enabled: true
  ssl.certificate_authorities: ["/usr/share/filebeat/certs/ca.crt"]
  ssl.verification_mode: full

logging.level: warning
""",
    )

    doc.add_heading("4.7 Creer Logstash", level=2)
    add_para(doc, "Fichier: deceptr\\elk\\logstash\\pipeline\\cowrie.conf. Il doit recevoir Beats TLS, parser JSON, indexer cowrie-* et pousser vers Redis.")
    add_code(
        doc,
        r"""
input {
  beats {
    port => 5044
    ssl_enabled => true
    ssl_certificate => "/usr/share/logstash/config/certs/logstash.crt"
    ssl_key => "/usr/share/logstash/config/certs/logstash.key"
  }
}

filter {
  if [eventid] {
    mutate { add_field => { "deceptr_pipeline" => "cowrie" } }
  }
}

output {
  elasticsearch {
    hosts => ["http://elasticsearch:9200"]
    user => "elastic"
    password => "${ELASTIC_PASSWORD}"
    index => "cowrie-%{+YYYY.MM}"
  }
  redis {
    host => "redis"
    port => 6379
    password => "${REDIS_PASSWORD}"
    data_type => "list"
    key => "${REDIS_QUEUE}"
  }
}
""",
    )

    doc.add_heading("4.8 Creer MySQL schema.sql", level=2)
    add_para(doc, "Le schema minimal doit contenir users, alerts, iocs, attackers et campaigns.")
    add_code(
        doc,
        r"""
CREATE DATABASE IF NOT EXISTS deceptr;
USE deceptr;

CREATE TABLE IF NOT EXISTS users (
  id INT AUTO_INCREMENT PRIMARY KEY,
  username VARCHAR(80) UNIQUE NOT NULL,
  password_hash VARCHAR(255) NOT NULL,
  role ENUM('admin','soc_analyst','threat_hunter','auditeur') DEFAULT 'soc_analyst',
  active TINYINT(1) DEFAULT 1,
  last_login DATETIME NULL
);

CREATE TABLE IF NOT EXISTS alerts (
  id INT AUTO_INCREMENT PRIMARY KEY,
  event_id VARCHAR(128),
  source_ip VARCHAR(64),
  severity VARCHAR(32),
  risk_score INT,
  mitre_technique VARCHAR(64),
  description TEXT,
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE IF NOT EXISTS iocs (
  id INT AUTO_INCREMENT PRIMARY KEY,
  indicator VARCHAR(255),
  type VARCHAR(64),
  reputation VARCHAR(64),
  first_seen DATETIME DEFAULT CURRENT_TIMESTAMP,
  last_seen DATETIME DEFAULT CURRENT_TIMESTAMP
);
""",
    )

    doc.add_heading("4.9 Creer le pipeline Python", level=2)
    add_para(doc, "Le pipeline peut etre minimal au debut, puis enrichi. Les modules obligatoires sont:")
    add_table(
        doc,
        ["Module", "Responsabilite minimale"],
        [
            ["collector.py", "Lire Redis et retourner les evenements"],
            ["normalizer.py", "Transformer Cowrie en champs standards: source.ip, username, password, event.type"],
            ["enricher.py", "Ajouter geoip/reputation si les cles existent"],
            ["correlator.py", "Regrouper par IP et campagne"],
            ["risk_scorer.py", "Calculer score et severite"],
            ["detector.py", "Creer alerte si login_attempt ou brute force"],
            ["storage.py", "Ecrire dans Elasticsearch, MySQL et MinIO"],
            ["main.py", "Boucle principale"],
            ["api/main.py", "FastAPI health, auth, alerts, stats"],
        ],
        widths=[1.65, 4.85],
    )
    add_para(doc, "Dockerfile minimal du pipeline/API:")
    add_code(
        doc,
        r"""
FROM python:3.11-slim
WORKDIR /app
RUN apt-get update && apt-get install -y --no-install-recommends gcc curl && rm -rf /var/lib/apt/lists/*
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY . .
RUN useradd -r -u 1001 -g root deceptr && chown -R deceptr:root /app
USER deceptr
""",
    )
    add_para(doc, "requirements.txt minimal:")
    add_code(
        doc,
        r"""
fastapi
uvicorn[standard]
elasticsearch
redis
mysql-connector-python
python-jose[cryptography]
passlib[bcrypt]
requests
minio
geoip2
""",
    )

    doc.add_heading("4.10 Creer dashboard et API", level=2)
    numbered(
        doc,
        [
            "Creer api/main.py avec /health, /api/auth/login, /api/alerts, /api/alerts/stats.",
            "Creer dashboard/index.html qui appelle http://127.0.0.1:8000.",
            "Verifier que le formulaire login envoie username/password vers /api/auth/login.",
            "Afficher KPIs, timeline, alertes recentes et IoC.",
        ],
    )

    doc.add_heading("4.11 Creer start.ps1 et stop.ps1", level=2)
    add_code(
        doc,
        r"""
# start.ps1
$Root = Split-Path -Parent $MyInvocation.MyCommand.Path
$DeceptrRoot = Join-Path $Root "deceptr"
$TpotRoot = Join-Path $Root "tpot-runtime"
Copy-Item -Force "$DeceptrRoot\tpot\filebeat\tpot-cowrie-to-deceptr.yml" "$TpotRoot\deceptr\filebeat\tpot-cowrie-to-deceptr.yml"
Copy-Item -Force "$DeceptrRoot\elk\certs\ca.crt" "$TpotRoot\deceptr\certs\ca.crt"
Push-Location $DeceptrRoot
docker compose -f docker-compose.yml -f docker-compose.tpot.yml up -d --build
Pop-Location
Push-Location $TpotRoot
docker compose up -d cowrie deceptr-tpot-forwarder
Pop-Location
""",
    )
    add_code(
        doc,
        r"""
# stop.ps1
Push-Location .\tpot-runtime
docker compose stop deceptr-tpot-forwarder cowrie
Pop-Location
Push-Location .\deceptr
docker compose -f docker-compose.yml -f docker-compose.tpot.yml stop
Pop-Location
""",
    )

    doc.add_heading("4.12 Validation d'un build zero", level=2)
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FROM-ZERO
powershell -ExecutionPolicy Bypass -File .\start.ps1
docker ps --format "table {{.Names}}\t{{.Status}}\t{{.Ports}}"
docker exec deceptr-tpot-forwarder filebeat test output -e --strict.perms=false
curl.exe -s -o NUL -w "%{http_code}" http://127.0.0.1:8000/health
curl.exe -s -o NUL -w "%{http_code}" http://127.0.0.1:8088/index.html
""",
    )
    add_callout(
        doc,
        "Critere de reussite",
        "Un build zero est valide quand Cowrie tourne sur 22/23, Filebeat annonce TLSv1.3 OK, l'API retourne 200, le dashboard retourne 200, Kibana repond, et un evenement test arrive dans cowrie-* puis deceptr-events-*.",
        fill=LIGHT_BLUE,
    )


def section_config(doc):
    doc.add_heading("5. Configuration manuelle des composants", level=1)
    doc.add_heading("5.1 Fichier .env principal", level=2)
    add_para(doc, "Chemin: D:\\assir\\Ismagi\\PFA\\DECEPTR-FINAL\\deceptr\\.env")
    add_table(
        doc,
        ["Variable", "Role", "Valeur lab actuelle / remarque"],
        [
            ["ELASTIC_PASSWORD", "Mot de passe elastic", "deceptr2025"],
            ["KIBANA_SYSTEM_PASS", "Compte technique Kibana", "deceptr-kibana"],
            ["REDIS_PASSWORD", "Protection Redis", "deceptr-redis-2026"],
            ["MYSQL_ROOT_PASSWORD", "Root MySQL", "root2025"],
            ["MYSQL_PASSWORD", "Utilisateur deceptr MySQL", "mysql2025"],
            ["MINIO_ROOT_USER", "Compte console MinIO", "deceptr_admin"],
            ["MINIO_ROOT_PASSWORD", "Mot de passe MinIO", "MotDePasseMinIO2026!"],
            ["JWT_SECRET", "Signature tokens API", "Changer en production"],
            ["VIRUSTOTAL_KEY / ABUSEIPDB_KEY", "Threat intelligence", "Optionnel; garder vide si pas de cle"],
            ["ALERT_TO", "Email destinataire alertes", "security@example.com par defaut"],
        ],
        widths=[1.8, 2.0, 2.7],
    )
    add_callout(doc, "Ne pas publier le .env", "Le fichier .env peut contenir des secrets. Il doit rester local au lab ou etre remplace par un gestionnaire de secrets en production.", fill=WARN)

    doc.add_heading("5.2 Certificats TLS Filebeat -> Logstash", level=2)
    add_para(doc, "Le projet contient deja les certificats dans deceptr\\elk\\certs. Pour regenerer manuellement:")
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr\elk\certs
openssl genrsa -out ca.key 4096
openssl req -x509 -new -nodes -key ca.key -sha256 -days 3650 -out ca.crt -subj "/CN=DECEPTR-CA"
openssl genrsa -out logstash.key 2048
openssl req -new -key logstash.key -out logstash.csr -subj "/CN=logstash"
openssl x509 -req -in logstash.csr -CA ca.crt -CAkey ca.key -CAcreateserial -out logstash.crt -days 825 -sha256
""",
    )
    add_para(doc, "Copier ensuite la CA vers le runtime T-Pot:")
    add_code(doc, r"copy .\ca.crt D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\deceptr\certs\ca.crt")

    doc.add_heading("5.3 T-Pot Cowrie et Filebeat", level=2)
    add_para(doc, "Chemin: tpot-runtime\\docker-compose.yml. Ce runtime est volontairement strict: seulement Cowrie officiel T-Pot et Filebeat forwarder.")
    add_table(
        doc,
        ["Service", "Configuration critique"],
        [
            ["cowrie", "Image ghcr.io/telekom-security/cowrie:24.04.1, ports 22:22 et 23:23, logs dans tpot-runtime\\data\\cowrie\\log"],
            ["deceptr-tpot-forwarder", "Image elastic/filebeat:8.13.0, lit /tpot/cowrie/log/cowrie.json, envoie vers logstash:5044 en TLS"],
        ],
        widths=[1.7, 4.8],
    )

    doc.add_heading("5.4 Logstash", level=2)
    add_para(doc, "Logstash ecoute sur TCP/5044 en Beats TLS. Il indexe les logs bruts dans cowrie-YYYY.MM et pousse une copie dans Redis.")
    add_code(
        doc,
        r"""
docker logs deceptr-logstash --tail 80
docker exec deceptr-tpot-forwarder filebeat test output -e --strict.perms=false
""",
    )
    doc.add_heading("5.5 Pipeline Python", level=2)
    add_table(
        doc,
        ["Fichier", "Role"],
        [
            ["collector.py", "Consomme Redis puis fallback Elasticsearch cowrie-*"],
            ["normalizer.py", "Transforme Cowrie brut en schema DECEPTR unifie"],
            ["enricher.py", "Ajoute GeoIP, VirusTotal, AbuseIPDB, Feodo"],
            ["correlator.py", "Regroupe evenements et IoC"],
            ["risk_scorer.py", "Calcule le score 0-100"],
            ["detector.py", "Declenche les regles de detection"],
            ["alerter.py", "Genere alertes et notifications"],
            ["storage.py", "Ecrit Elasticsearch, MySQL et MinIO"],
            ["api/main.py", "Expose FastAPI et les endpoints dashboard"],
        ],
        widths=[1.75, 4.75],
    )
    doc.add_heading("5.6 Comptes par defaut", level=2)
    add_table(
        doc,
        ["Interface", "URL", "Utilisateur", "Mot de passe lab"],
        [
            ["Dashboard DECEPTR", "http://127.0.0.1:8088/index.html?v=3", "admin", "deceptr2025"],
            ["Kibana", "http://127.0.0.1:5601", "elastic", "deceptr2025"],
            ["API Docs", "http://127.0.0.1:8000/docs", "JWT via /api/auth/login", "admin / deceptr2025"],
            ["MinIO", "http://127.0.0.1:9001", "deceptr_admin", "MotDePasseMinIO2026!"],
        ],
        widths=[1.45, 2.25, 1.25, 1.55],
    )


def section_start(doc):
    doc.add_heading("6. Demarrage manuel de toute l'architecture", level=1)
    doc.add_heading("6.1 Methode recommandee", level=2)
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL
powershell -ExecutionPolicy Bypass -File .\start.ps1
""",
    )
    doc.add_heading("6.2 Methode totalement manuelle", level=2)
    add_para(doc, "Etape 1 - preparer les fichiers Filebeat et certificats:")
    add_code(
        doc,
        r"""
mkdir D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\deceptr\filebeat
mkdir D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\deceptr\certs
copy D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr\tpot\filebeat\tpot-cowrie-to-deceptr.yml D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\deceptr\filebeat\
copy D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr\elk\certs\ca.crt D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\deceptr\certs\
mkdir D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime\data\cowrie\log
""",
    )
    add_para(doc, "Etape 2 - demarrer DECEPTR sans Cowrie local:")
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr
docker compose -f docker-compose.yml -f docker-compose.tpot.yml up -d --build
""",
    )
    add_para(doc, "Etape 3 - demarrer Cowrie T-Pot et le forwarder:")
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime
docker compose up -d cowrie deceptr-tpot-forwarder
""",
    )
    doc.add_heading("6.3 Verifier l'etat", level=2)
    add_code(
        doc,
        r"""
docker ps --format "table {{.Names}}\t{{.Status}}\t{{.Ports}}"
curl.exe -s -o NUL -w "%{http_code}" http://127.0.0.1:8088/index.html?v=3
curl.exe -s -o NUL -w "%{http_code}" http://127.0.0.1:8000/health
curl.exe -s -o NUL -w "%{http_code}" http://127.0.0.1:5601
docker exec deceptr-tpot-forwarder filebeat test output -e --strict.perms=false
""",
    )
    add_table(
        doc,
        ["Test", "Resultat attendu"],
        [
            ["Dashboard", "HTTP 200"],
            ["API health", "HTTP 200"],
            ["Kibana", "HTTP 302 ou page login"],
            ["Filebeat test output", "TLS handshake OK, TLSv1.3"],
            ["Cowrie", "Ports 22 et 23 publies"],
        ],
        widths=[2.0, 4.5],
    )


def section_kibana(doc):
    doc.add_heading("7. Configuration manuelle de Kibana et des dashboards", level=1)
    numbered(
        doc,
        [
            "Ouvrir http://127.0.0.1:5601.",
            "Se connecter avec elastic / deceptr2025.",
            "Aller dans Stack Management > Data Views.",
            "Creer le data view cowrie-* avec le champ temps @timestamp.",
            "Creer le data view deceptr-events-* avec le champ temps @timestamp.",
            "Aller dans Discover et verifier que les evenements apparaissent.",
        ],
    )
    add_table(
        doc,
        ["Dashboard Kibana conseille", "Visualisation", "Source"],
        [
            ["Activite Cowrie", "Histogramme par heure", "cowrie-*"],
            ["Top IP attaquantes", "Table terms source.ip", "deceptr-events-*"],
            ["Identifiants testes", "Table username/password", "deceptr-events-*"],
            ["MITRE ATT&CK", "Bar chart mitre.technique", "deceptr-events-*"],
            ["Severite", "Donut severity", "deceptr-events-*"],
            ["Carte GeoIP", "Map source.geo.location", "deceptr-events-*"],
        ],
        widths=[1.8, 2.3, 2.4],
    )
    add_callout(doc, "Pourquoi Kibana peut sembler vide", "Kibana n'invente pas les donnees. Il faut au moins un evenement Cowrie traite dans cowrie-* ou deceptr-events-* avant que Discover et Lens affichent quelque chose.", fill=WARN)


def section_tests(doc):
    doc.add_heading("8. Tests utilisateur et validation bout-en-bout", level=1)
    doc.add_heading("8.1 Test interactif SSH/Telnet", level=2)
    add_code(
        doc,
        r"""
ssh root@127.0.0.1 -p 22
# Essayer quelques mots de passe: admin, 123456, password

telnet 127.0.0.1 23
# Essayer login: admin / password: admin
""",
    )
    add_para(doc, "Ces actions doivent creer des lignes dans tpot-runtime\\data\\cowrie\\log\\cowrie.json.")
    doc.add_heading("8.2 Test E2E automatise", level=2)
    add_code(
        doc,
        r"""
powershell -ExecutionPolicy Bypass -File D:\assir\Ismagi\PFA\DECEPTR-FINAL\deceptr\scripts\e2e-smoke.ps1 -TpotRoot D:\assir\Ismagi\PFA\DECEPTR-FINAL\tpot-runtime -WaitSeconds 75
""",
    )
    add_table(
        doc,
        ["Champ attendu", "Exemple valide"],
        [
            ["status", "OK"],
            ["tpot_cowrie", "running"],
            ["tpot_forwarder_tls", "TLSv1.3"],
            ["raw_index", "cowrie-2026.06"],
            ["enriched_index", "deceptr-events-2026.06"],
            ["enriched_type", "login_attempt"],
            ["mitre", "T1110"],
        ],
        widths=[2.2, 4.3],
    )
    doc.add_heading("8.3 Verification API", level=2)
    add_code(
        doc,
        r"""
$body = @{username='admin'; password='deceptr2025'} | ConvertTo-Json
$token = (Invoke-RestMethod -Method Post -Uri http://127.0.0.1:8000/api/auth/login -ContentType 'application/json' -Body $body).access_token
Invoke-RestMethod -Uri http://127.0.0.1:8000/api/alerts -Headers @{Authorization="Bearer $token"}
Invoke-RestMethod -Uri http://127.0.0.1:8000/api/alerts/stats -Headers @{Authorization="Bearer $token"}
""",
    )


def section_scripts(doc):
    doc.add_heading("9. Scripts d'automatisation du projet", level=1)
    add_table(
        doc,
        ["Script", "Emplacement", "Role"],
        [
            ["start.ps1", "DECEPTR-FINAL", "Prepare fichiers, demarre DECEPTR, Cowrie T-Pot et forwarder"],
            ["stop.ps1", "DECEPTR-FINAL", "Arrete les services du projet sans supprimer les volumes"],
            ["e2e-smoke.ps1", "deceptr\\scripts", "Injecte un evenement et verifie toute la chaine"],
            ["start-architecture.ps1", "deceptr\\scripts", "Redirection vers le start final"],
            ["stop-architecture.ps1", "deceptr\\scripts", "Redirection vers le stop final"],
        ],
        widths=[1.55, 1.75, 3.2],
    )
    doc.add_heading("9.1 Script start.ps1", level=2)
    add_para(doc, "Le script fait trois choses: verifie Docker, copie la configuration Filebeat/certificats vers tpot-runtime, puis lance les deux compose.")
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL
powershell -ExecutionPolicy Bypass -File .\start.ps1
""",
    )
    doc.add_heading("9.2 Script stop.ps1", level=2)
    add_code(
        doc,
        r"""
cd D:\assir\Ismagi\PFA\DECEPTR-FINAL
powershell -ExecutionPolicy Bypass -File .\stop.ps1
""",
    )


def section_ops(doc):
    doc.add_heading("10. Exploitation SOC et rapport final", level=1)
    add_para(doc, "Le SOC exploite DECEPTR via trois niveaux: dashboard rapide, Kibana pour investigation, API pour extraction et integration.")
    add_table(
        doc,
        ["Besoin SOC", "Outil", "Action"],
        [
            ["Voir le niveau de menace", "Dashboard DECEPTR", "Ouvrir /index.html?v=3 et verifier KPIs"],
            ["Analyser une IP", "API / Kibana", "Filtrer source.ip puis voir credentials, pays, MITRE"],
            ["Voir brute force", "Kibana", "Filtrer event.type=login_attempt et MITRE T1110"],
            ["Exporter un rapport", "API DGSSI", "Consulter /api/dgssi/rapport ou /api/dgssi/rapport/json"],
            ["Traiter une alerte", "Dashboard/API", "Acknowledge et documenter la decision"],
        ],
        widths=[1.65, 1.65, 3.2],
    )
    doc.add_heading("10.1 Structure du rapport final", level=2)
    add_table(
        doc,
        ["Section", "Contenu"],
        [
            ["Resume executif", "Nombre d'alertes, severite dominante, tendance 24h, decision a prendre"],
            ["Architecture", "Schema reseau, fonctionnel et technique"],
            ["Sources de donnees", "Cowrie, Filebeat, Logstash, Redis, pipeline"],
            ["Alertes", "Liste par severite, IP, MITRE, score"],
            ["IoC", "IP, reputation, geolocalisation, historique"],
            ["Timeline", "Evenements par heure et par campagne"],
            ["Recommandations", "Court terme, moyen terme, durcissement"],
            ["Annexes", "Commandes, hashes, versions, captures dashboard"],
        ],
        widths=[1.8, 4.7],
    )


def section_maintenance(doc):
    doc.add_heading("11. Maintenance, sauvegarde et nettoyage Docker", level=1)
    doc.add_heading("11.1 Sauvegarder les donnees", level=2)
    add_code(
        doc,
        r"""
docker run --rm -v deceptr_elastic_data:/data -v D:\backups:/backup alpine tar czf /backup/elastic_data.tgz /data
docker run --rm -v deceptr_mysql_data:/data -v D:\backups:/backup alpine tar czf /backup/mysql_data.tgz /data
docker run --rm -v deceptr_minio_data:/data -v D:\backups:/backup alpine tar czf /backup/minio_data.tgz /data
""",
    )
    doc.add_heading("11.2 Nettoyer uniquement DECEPTR", level=2)
    add_para(doc, "Ne jamais lancer docker system prune sans comprendre l'impact. Pour ce projet, supprimer seulement les conteneurs nommes DECEPTR si necessaire:")
    add_code(
        doc,
        r"""
docker rm -f cowrie deceptr-tpot-forwarder deceptr-logstash deceptr-redis deceptr-pipeline deceptr-elasticsearch deceptr-mysql deceptr-minio deceptr-kibana deceptr-api deceptr-dashboard deceptr-elastalert
""",
    )
    add_callout(doc, "Attention", "Ne pas supprimer les projets non lies comme Pentagi, SIEM, Mongo ou autres environnements Docker personnels.", fill=WARN)
    doc.add_heading("11.3 Logs utiles", level=2)
    add_code(
        doc,
        r"""
docker logs cowrie --tail 80
docker logs deceptr-tpot-forwarder --tail 80
docker logs deceptr-logstash --tail 120
docker logs deceptr-pipeline --tail 120
docker logs deceptr-api --tail 120
docker logs deceptr-kibana --tail 120
""",
    )


def section_checklist(doc):
    doc.add_heading("12. Checklist complete", level=1)
    add_table(
        doc,
        ["Phase", "Point de controle", "OK"],
        [
            ["Preparation", "Docker Desktop demarre et docker version repond", "[ ]"],
            ["Preparation", "Ports 22, 23, 5044, 5601, 8000, 8088, 9000, 9001, 9200 libres", "[ ]"],
            ["Preparation", "Dossier final present dans D:\\assir\\Ismagi\\PFA\\DECEPTR-FINAL", "[ ]"],
            ["Configuration", ".env rempli avec mots de passe lab ou production", "[ ]"],
            ["Configuration", "ca.crt copie vers tpot-runtime\\deceptr\\certs", "[ ]"],
            ["Demarrage", "start.ps1 termine sans erreur", "[ ]"],
            ["Conteneurs", "12 conteneurs attendus visibles dans docker ps", "[ ]"],
            ["TLS", "filebeat test output indique TLSv1.3 OK", "[ ]"],
            ["Dashboard", "http://127.0.0.1:8088/index.html?v=3 retourne 200", "[ ]"],
            ["API", "http://127.0.0.1:8000/health retourne 200", "[ ]"],
            ["Kibana", "http://127.0.0.1:5601 retourne 302 ou login", "[ ]"],
            ["Kibana", "Data views cowrie-* et deceptr-events-* crees", "[ ]"],
            ["Honeypot", "Tentative SSH ou Telnet genere cowrie.json", "[ ]"],
            ["Pipeline", "e2e-smoke.ps1 retourne status OK", "[ ]"],
            ["Rapport", "Captures ou exports disponibles pour le rapport PFA", "[ ]"],
        ],
        widths=[1.25, 4.55, 0.7],
    )


def section_conclusion(doc):
    doc.add_heading("13. Conclusion", level=1)
    add_para(doc, "DECEPTR v1 MVP repond a l'objectif du projet: deployer une plateforme de cyberdeception reproductible, dockerisee, basee sur Cowrie T-Pot et un pipeline complet de collecte, enrichissement, detection, stockage, visualisation et alerting.")
    add_table(
        doc,
        ["Objectif", "Resultat"],
        [
            ["Honeypot reel", "Cowrie officiel T-Pot expose sur SSH 22 et Telnet 23"],
            ["Collecte securisee", "Filebeat vers Logstash en TLS 1.3"],
            ["Pipeline complet", "Redis + Python normalisation/enrichissement/correlation/MITRE/risk/alertes"],
            ["Stockage", "Elasticsearch, MySQL et MinIO"],
            ["Visualisation", "Kibana, Dashboard DECEPTR et API FastAPI"],
            ["Alerting", "ElastAlert 2 et alerter Python"],
            ["Reproductibilite", "Un seul dossier final avec start.ps1 et stop.ps1"],
        ],
        widths=[1.8, 4.7],
    )
    add_callout(doc, "Projet pret", "La partie logicielle correspond au schema MVP. La partie restante pour une production reelle est l'infrastructure physique: firewall pfSense/FortiGate, segmentation DMZ/LAN/SOC, VPN/MFA, certificats publics et politique de sauvegarde.", fill=LIGHT_BLUE)


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_static_toc(doc)
    section_errors(doc)
    section_architecture(doc)
    section_prereq(doc)
    section_build_zero(doc)
    section_install(doc)
    section_config(doc)
    section_start(doc)
    section_kibana(doc)
    section_tests(doc)
    section_scripts(doc)
    section_ops(doc)
    section_maintenance(doc)
    section_checklist(doc)
    section_conclusion(doc)
    doc.core_properties.title = "DECEPTR v1 MVP - Guide complet installation et configuration"
    doc.core_properties.subject = "Cyberdeception, T-Pot Cowrie, ELK, Docker"
    doc.core_properties.author = "DECEPTR Project"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
